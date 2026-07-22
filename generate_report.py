import os
import subprocess
import sys
import re
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def get_packet_sizes(student_id):
    mid = len(student_id) // 2
    h1 = student_id[:mid]
    h2 = student_id[mid:]
    h1_rev = h1[::-1]
    h2_rev = h2[::-1]
    return [int(h1), int(h2), int(h1_rev), int(h2_rev)]

def run_simulation(packet_size):
    # Runs the ns-3 simulation and returns the full output text and the throughput value
    cmd = ["./ns3", "run", f"scratch/first.cc --packetSize={packet_size}"]
    result = subprocess.run(cmd, capture_output=True, text=True, cwd="ns-allinone-3.43/ns-3.43")
    full_output = result.stdout + "\n" + result.stderr
    
    # Filter the output to keep only the relevant lines (At time ... and FlowID ...)
    lines = full_output.split('\n')
    relevant_lines = []
    capture = False
    for line in lines:
        if "At time" in line or "TraceDelay" in line or "FlowID:" in line or "Tx Bytes" in line or "Rx Bytes" in line or "Mean Delay" in line or "Throughput" in line:
            relevant_lines.append(line)
            
    filtered_text = "\n".join(relevant_lines)
    
    match = re.search(r"Throughput:\s*([\d\.]+)\s*bps", full_output)
    throughput = float(match.group(1)) if match else 0.0
    return filtered_text, throughput

def text_to_image(text, filename):
    # Renders text as a premium terminal-style screenshot
    fig, ax = plt.subplots(figsize=(8.5, 3.5))
    ax.text(0.02, 0.95, text, family='monospace', fontsize=9.5, color='#A8FF60',
            verticalalignment='top', horizontalalignment='left')
    ax.set_facecolor('#1A1A24')
    fig.patch.set_facecolor('#1A1A24')
    ax.axis('off')
    plt.savefig(filename, facecolor=fig.get_facecolor(), edgecolor='none', bbox_inches='tight', dpi=200)
    plt.close()

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    # Sets margins (padding) for a table cell
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy
        run.font.name = 'Georgia'
        # Add a thin bottom border or underline using paragraph formatting if desired, but bold navy Georgia is premium
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x4A, 0x77, 0x7A) # Slate Teal
        run.font.name = 'Georgia'
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.name = 'Arial'
    return p

def main():
    student_id = "24252627"
    if len(sys.argv) > 1:
        student_id = sys.argv[1]
        
    sizes = get_packet_sizes(student_id)
    
    results = {}
    throughputs = []
    
    # Create screenshots directory
    os.makedirs("scratch_screenshots", exist_ok=True)
    
    for size in sizes:
        txt, tp = run_simulation(size)
        results[size] = {
            "text": txt,
            "throughput": tp
        }
        throughputs.append(tp)
        
        # Render screenshot image
        img_path = f"scratch_screenshots/screenshot_{size}.png"
        text_to_image(txt, img_path)
        results[size]["img"] = img_path

    # Generate the Throughput vs Packet Size plot
    plt.figure(figsize=(7.5, 5))
    paired = sorted(zip(sizes, throughputs))
    sorted_sizes, sorted_throughputs = zip(*paired)
    
    plt.plot(sorted_sizes, sorted_throughputs, marker='o', linestyle='-', color='#1B365D', linewidth=2.5, markersize=8)
    
    # Add labels
    for s, t in zip(sorted_sizes, sorted_throughputs):
        plt.annotate(f"{t:.2f} bps", (s, t), textcoords="offset points", xytext=(0,10), ha='center', fontsize=9, fontweight='bold', color='#333333')
        
    plt.title("Throughput vs. Packet Size", fontsize=13, fontweight='bold', pad=15, color='#1B365D')
    plt.xlabel("Packet Size (bytes)", fontsize=11, labelpad=8)
    plt.ylabel("Throughput (bps)", fontsize=11, labelpad=8)
    plt.grid(True, linestyle='--', alpha=0.5)
    
    plot_path = "scratch_screenshots/throughput_plot.png"
    plt.savefig(plot_path, dpi=300, bbox_inches='tight')
    plt.close()

    # Build the Word Document
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("LAB REPORT")
    title_run.font.size = Pt(26)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    title_run.font.name = 'Georgia'
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    sub_run = subtitle_p.add_run("Task 3: Introduction to Network Simulator (ns-3)")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub_run.font.name = 'Georgia'
    
    # Metadata Table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata = [
        ("Course Task:", "Task 3 - Introduction to network simulator (ns3)"),
        ("Student ID:", student_id),
        ("Date:", "July 21, 2026"),
        ("Simulator Version:", "ns-allinone-3.43 (ns-3.43)")
    ]
    for idx, (label, val) in enumerate(metadata):
        row = table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        # Style label
        lbl_p = cell_lbl.paragraphs[0]
        lbl_p.paragraph_format.space_after = Pt(2)
        r_lbl = lbl_p.add_run(label)
        r_lbl.bold = True
        r_lbl.font.size = Pt(10.5)
        r_lbl.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r_lbl.font.name = 'Arial'
        
        # Style value
        val_p = cell_val.paragraphs[0]
        val_p.paragraph_format.space_after = Pt(2)
        r_val = val_p.add_run(val)
        r_val.font.size = Pt(10.5)
        r_val.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        r_val.font.name = 'Arial'
        
        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)
        
    doc.add_page_break()
    
    # Section 1: Introduction & Edits
    add_styled_heading(doc, "1. Introduction & Modifications to first.cc", level=1)
    
    intro_text = (
        "In this lab, we explored the foundational components of the ns-3 network simulator. "
        "The topology consists of a simple point-to-point network between two nodes (n0 and n1), "
        "configured with a 5 Mbps data rate and a 2 ms propagation delay. Node 0 acts as a UDP echo client, "
        "and Node 1 acts as a UDP echo server. To enhance the simulation, we modified the default tutorial "
        "script first.cc to include NetAnim visualization, PCAP tracing, and FlowMonitor statistics, "
        "and made the client's packet size configurable via command line arguments."
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(intro_text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    
    add_styled_heading(doc, "Modifications Summary:", level=2)
    mod_points = [
        ("NetAnim Integration: ", "Added the #include \"ns3/netanim-module.h\" header and instantiated AnimationInterface anim(\"first.xml\") before the simulation run. This creates an XML animation trace file that can be loaded in NetAnim for visualization."),
        ("PCAP Trait Capture: ", "Added pointToPoint.EnablePcapAll(\"first\") to capture raw packet logs (.pcap format) for all interfaces, enabling offline packet inspection in Wireshark."),
        ("FlowMonitor Statistics: ", "Added the #include \"ns3/flow-monitor-module.h\" header, installed the FlowMonitorHelper on all nodes, and processed flow statistics (Tx/Rx bytes, delay, throughput) at the end of the simulation."),
        ("Dynamic Packet Size: ", "Added a command-line parameter \"packetSize\" via the CommandLine helper, allowing the packet size to be dynamically configured during invocation without re-compiling the code.")
    ]
    
    for bold_prefix, desc in mod_points:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.paragraph_format.space_after = Pt(4)
        run_bold = bullet.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(10.5)
        run_bold.font.name = 'Arial'
        run_desc = bullet.add_run(desc)
        run_desc.font.size = Pt(10.5)
        run_desc.font.name = 'Arial'
        
    doc.add_page_break()
    
    # Section 2: Screenshots & Outputs
    add_styled_heading(doc, "2. Simulation Outputs for Different Packet Sizes", level=1)
    
    for size in sizes:
        add_styled_heading(doc, f"Packet Size: {size} Bytes", level=2)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"Below is the programmatic screenshot of the terminal log output for a packet payload size of {size} bytes:")
        r.font.size = Pt(11)
        r.font.name = 'Arial'
        r.italic = True
        
        # Add screenshot image
        doc.add_picture(results[size]["img"], width=Inches(6.0))
        
        # Add caption
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(4)
        cap_p.paragraph_format.space_after = Pt(14)
        cap_r = cap_p.add_run(f"Figure {sizes.index(size) + 1}: Simulation logs for {size} bytes")
        cap_r.font.size = Pt(9.5)
        cap_r.font.name = 'Arial'
        cap_r.italic = True
        cap_r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        
    doc.add_page_break()
    
    # Section 3: Plot
    add_styled_heading(doc, "3. Throughput vs. Packet Size Graph", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("The relationship between packet size (derived from the student ID) and the measured network throughput is plotted below:")
    r.font.size = Pt(11)
    r.font.name = 'Arial'
    
    doc.add_picture(plot_path, width=Inches(5.5))
    
    # Add caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(4)
    cap_p.paragraph_format.space_after = Pt(14)
    cap_r = cap_p.add_run("Figure 5: Plot of Throughput (bps) vs. Packet Size (bytes)")
    cap_r.font.size = Pt(9.5)
    cap_r.font.name = 'Arial'
    cap_r.italic = True
    cap_r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    
    # Section 4: Explanation
    add_styled_heading(doc, "4. Explanation of the Throughput vs. Packet Size Graph", level=1)
    
    explanation_paras = [
        "The graph shows a perfectly linear relationship between the packet size and the measured throughput. "
        "As the packet size increases, the throughput also increases proportionally. To understand this behavior, "
        "we must analyze how the simulation measures throughput.",
        
        "The client application in our script is configured to send exactly one packet (MaxPackets = 1). "
        "The simulation is configured to start client transmission at 2.0 seconds and stops at 20.0 seconds, "
        "which gives an active simulation window of exactly 18.0 seconds. The FlowMonitor helper calculates "
        "the throughput by taking the total received bits (Rx Bytes × 8) and dividing them by this 18-second "
        "simulation time window:",
        
        "Throughput (bps) = (Rx Bytes * 8.0) / 18.0",
        
        "Since only a single packet is transmitted during the entire duration of the simulation, the total received "
        "bytes (Rx Bytes) measured at the IP layer corresponds to exactly one packet. This includes the UDP payload "
        "size (P), the UDP header (8 bytes), and the IPv4 header (20 bytes). Thus, the Rx Bytes is always: "
        "Rx Bytes = P + 28 bytes.",
        
        "Substituting this back into the formula yields a linear equation with respect to payload size P:",
        
        "Throughput (bps) = ( (P + 28) * 8.0 ) / 18.0 = (8/18) * P + (224/18)",
        
        "Since the coefficient (8/18) is positive, the throughput increases strictly linearly with the packet size P. "
        "If the simulation were configured to transmit packets continuously (i.e., saturating the channel), the throughput "
        "would eventually level off and be bounded by the point-to-point link capacity (5 Mbps). However, because of "
        "the single-packet configuration, the measured throughput is simply a linear function of the packet size."
    ]
    
    for para in explanation_paras:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(10)
        
        # Center formula paragraphs
        if "Throughput (bps)" in para or "Throughput (bps) =" in para:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(para)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Courier New'
        else:
            run = p.add_run(para)
            run.font.size = Pt(11)
            run.font.name = 'Arial'

    # Save document
    doc_name = "Task 3 - Introduction to network simulator (ns3) Report.docx"
    doc.save(doc_name)
    
    # Save a copy as Report.docx for convenience
    doc.save("Report.docx")
    
    print(f"Report successfully saved as: {os.path.abspath(doc_name)}")
    print(f"Convenience copy saved as: {os.path.abspath('Report.docx')}")

if __name__ == "__main__":
    main()
